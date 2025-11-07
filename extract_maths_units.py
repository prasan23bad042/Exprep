"""
Script to extract content from 9th Maths unit documents
and create a TypeScript data file
"""

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    import json
    import re
    from html import escape
    
    def extract_table_html(table):
        """Convert a Word table to HTML"""
        html = '<table class="w-full border-collapse border border-gray-300 my-4">\n'
        
        for i, row in enumerate(table.rows):
            html += '  <tr>\n'
            for cell in row.cells:
                tag = 'th' if i == 0 else 'td'
                cell_class = 'border border-gray-300 px-4 py-2 font-semibold bg-gray-100' if i == 0 else 'border border-gray-300 px-4 py-2'
                html += f'    <{tag} class="{cell_class}">{escape(cell.text.strip())}</{tag}>\n'
            html += '  </tr>\n'
        
        html += '</table>\n'
        return html
    
    def process_paragraph(para):
        """Process a paragraph and return HTML"""
        text = para.text.strip()
        if not text:
            return ''
        
        # Check if it's a heading
        if para.style.name.startswith('Heading'):
            level_match = re.search(r'\d+', para.style.name)
            if level_match:
                level = min(int(level_match.group()), 6)  # Max h6
                return f'<h{level} class="font-bold text-lg mt-4 mb-2">{escape(text)}</h{level}>\n'
        
        # Build HTML with formatting
        html_parts = []
        for run in para.runs:
            run_text = escape(run.text)
            if run.bold:
                run_text = f'<strong>{run_text}</strong>'
            if run.italic:
                run_text = f'<em>{run_text}</em>'
            if run.underline:
                run_text = f'<u>{run_text}</u>'
            html_parts.append(run_text)
        
        html_text = ''.join(html_parts) if html_parts else escape(text)
        
        # Check if it's a list item
        if para.style.name.startswith('List'):
            return f'<li>{html_text}</li>\n'
        
        return f'<p class="mb-2">{html_text}</p>\n'
    
    def extract_unit_content(doc_path):
        """Extract content from a unit document"""
        doc = Document(doc_path)
        content_html = []
        in_list = False
        
        for element in doc.element.body:
            if hasattr(element, 'tag'):
                if element.tag.endswith('tbl'):
                    # It's a table
                    if in_list:
                        content_html.append('</ul>\n')
                        in_list = False
                    
                    for table in doc.tables:
                        if table._element == element:
                            content_html.append(extract_table_html(table))
                            break
                
                elif element.tag.endswith('p'):
                    # It's a paragraph
                    for para in doc.paragraphs:
                        if para._element == element:
                            para_html = process_paragraph(para)
                            
                            if para_html.startswith('<li>'):
                                if not in_list:
                                    content_html.append('<ul class="list-disc list-inside ml-4 mb-2">\n')
                                    in_list = True
                                content_html.append(para_html)
                            else:
                                if in_list:
                                    content_html.append('</ul>\n')
                                    in_list = False
                                content_html.append(para_html)
                            break
        
        if in_list:
            content_html.append('</ul>\n')
        
        return ''.join(content_html)
    
    # Process all units
    units_data = {}
    
    for i in range(1, 11):
        unit_file = f'd:/Exprep/9TH MATHS/UNIT {i}.docx' if i <= 7 else f'd:/Exprep/9TH MATHS/Unit {i}.docx'
        
        try:
            print(f'Processing Unit {i}...')
            content = extract_unit_content(unit_file)
            units_data[f'unit-{i}'] = content
            print(f'Unit {i} processed successfully!')
        except Exception as e:
            print(f'Error processing Unit {i}: {str(e)}')
    
    # Save to JSON file for easier handling
    with open('d:/Exprep/maths_units_content.json', 'w', encoding='utf-8') as f:
        json.dump(units_data, f, ensure_ascii=False, indent=2)
    
    print('\nAll units processed! Content saved to maths_units_content.json')
    print('You can now use this to populate the TypeScript data file.')

except ImportError:
    print('Error: python-docx library not found.')
    print('Please install it using: pip install python-docx')
except Exception as e:
    print(f'Error: {str(e)}')
